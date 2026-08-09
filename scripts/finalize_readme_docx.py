from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


DOCUMENT_PATH = Path("readme.docx")
CONCURRENCY_IMAGE = Path("artifacts/concurrency-100-requests-with-rejections.png")

NAVY = "17365D"
BLUE = "1F4E78"
TEAL = "0F6B78"
LIGHT_BLUE = "DDEBF7"
LIGHT_GRAY = "F3F6F8"
MID_GRAY = "D9E2F3"
TEXT = RGBColor(38, 50, 56)


def set_east_asia_font(element, name):
    element.rPr.rFonts.set(qn("w:eastAsia"), name)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def add_bottom_border(paragraph, color=BLUE, size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def remove_paragraph(paragraph):
    paragraph._element.getparent().remove(paragraph._element)


def insert_before(anchor, text="", style=None):
    paragraph = anchor.insert_paragraph_before(text)
    if style:
        paragraph.style = style
    return paragraph


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("- ")
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.append(field_begin)
    run._r.append(instruction)
    run._r.append(field_end)
    paragraph.add_run(" -")


def ensure_style(document, name, style_type=WD_STYLE_TYPE.PARAGRAPH):
    try:
        return document.styles[name]
    except KeyError:
        return document.styles.add_style(name, style_type)


document = Document(DOCUMENT_PATH)
ensure_style(document, "Caption")

# Remove the accumulated change-log block. Its final state is summarized below.
paragraphs = document.paragraphs
start = next(i for i, p in enumerate(paragraphs) if p.text == "프로젝트 안정성 개선 사항")
end = next(i for i, p in enumerate(paragraphs[start:], start) if p.text.strip() == "swagger api 실행 화면")
for paragraph in paragraphs[start:end]:
    remove_paragraph(paragraph)

# Correct stale terminology and final verification counts.
replacements = {
    "• 공개 그룹·솔로 페이지는 artist_channel 테이블에 저장한다.":
        "• 공개 그룹·솔로 페이지는 channel 테이블에 저장한다.",
    "• ARTIST와 MANAGER의 채널 소속은 artist_channel_user 단일 중간 테이블에 저장한다.":
        "• ARTIST와 MANAGER의 채널 소속은 channel_user 단일 중간 테이블에 저장한다.",
    "• artist_channel_user는 channel_id와 user_id를 가지며 별도의 관계 역할 컬럼 없이 users.role로 ARTIST와 MANAGER를 구분한다.":
        "• channel_user는 channel_id와 user_id를 가지며 users.role로 ARTIST와 MANAGER를 구분한다.",
    "• 게시글, 상품, 공연, 팬 멤버십은 artist_id 대신 channel_id로 artist_channel을 참조한다.":
        "• 게시글, 상품, 공연, 팬 멤버십은 channel_id로 channel을 참조한다.",
    "• 기존 artist_member와 MANAGER 테이블은 제거하였다.":
        "• 이전 아티스트 전용 관계 모델은 제거하고 채널 중심 모델로 통합하였다.",
    "JUnit 5 기반 단위·통합 테스트 11개 클래스, 총 27건을 실행했으며 실패와 건너뜀 없이 모두 통과하였다.":
        "JUnit 5 기반 단위·통합 테스트 13개 클래스, 총 45건을 실행했으며 실패와 건너뜀 없이 모두 통과하였다.",
    " swagger api 실행 화면": "9. API 실행 화면",
    "이미 팔린 자색": "이미 예매된 좌석",
    "공연 예약": "공연 좌석 예매 성공",
}
for paragraph in document.paragraphs:
    if paragraph.text in replacements:
        paragraph.text = replacements[paragraph.text]

# Remove meaningless stray captions.
for paragraph in list(document.paragraphs):
    if paragraph.text.strip() in {".", "’"}:
        if paragraph._p.xpath(".//w:drawing"):
            for text_node in paragraph._p.xpath(".//w:t"):
                text_node.text = ""
        else:
            remove_paragraph(paragraph)

# Normalize paragraphs that were accidentally stored as headings.
normal_texts = {
    "채널: 그룹 또는 솔로 아티스트를 나타내는 공개 페이지이자 게시글·상품·공연·팬 커뮤니티의 소속 단위이다.",
    "ARTIST 회원: users.role이 ARTIST인 로그인 계정이며 하나의 채널에만 소속된다.",
    "MANAGER 회원: users.role이 MANAGER인 로그인 계정이며 하나의 채널만 담당한다.",
    "댓글은 게시글에 직접 작성하는 최상위 댓글과 특정 댓글 또는 답글에 작성하는 답글로 구분한다.",
    "답글에 다시 답글을 작성할 수 있으나, 화면에서는 모든 답글을 동일한 깊이로 표시한다. 답글 대상은 별도로 저장하여 어떤 사용자 또는 댓글에 대한 답글인지 확인할 수 있도록 한다.",
    "공연 등록: 현재 시각 < 예매 시작 ≤ 예매 종료 ≤ 공연 일시 조건을 만족해야 한다.",
    "예매 시작 전: 공연의 모든 정보와 좌석 등록·수정·삭제가 가능하다.",
    "예매 시작 후: 공연 제목·장소 및 공연 상태만 변경할 수 있다.",
    "예매 시작 후: 공연 일시와 예매 시작·종료 일시는 변경할 수 없다.",
    "예매 시작 후: 좌석 개별·일괄 등록, 수정, 삭제는 모두 허용하지 않는다.",
    "공연 삭제 API는 제공하지 않는다.",
    "좌석 목록 조회는 MyBatis 동적 SQL로 필터링하며 page와 size로 페이징할 수 있다.",
    "선택 파라미터: section, grade, status, page(기본 0), size(기본 100, 최대 500)",
    "응답은 content와 page, size, totalElements, totalPages, first, last 페이지 메타데이터를 포함한다.",
    "좌석 일괄 등록의 startNumber는 선택값이다.",
    "생략 시 같은 공연·구역에 등록된 숫자 좌석 중 가장 큰 번호의 다음 번호부터 생성한다. 해당 구역에 좌석이 없으면 1번부터 생성한다.",
    "중간의 빈 번호는 자동으로 채우지 않으며 필요하면 startNumber를 직접 지정한다.",
    "응답은 section, 실제 startNumber, endNumber, createdCount와 생성된 seats 배열을 포함한다.",
    "자동 번호 계산과 등록은 공연 단위 비관적 잠금이 적용된 하나의 트랜잭션에서 처리한다.",
    "DELETE /api/seats/{seatId}",
    "담당 채널의 매니저가 좌석 ID로 좌석 하나를 삭제한다.",
    "응답: 204 No Content",
    "삭제 제한: AVAILABLE 상태이고 예매 이력이 없는 좌석만 삭제할 수 있으며, 담당 채널 매니저 권한이 필요하다.",
    "운영 정책: 좌석 삭제는 공연 예매 시작 전까지만 허용한다.",
}
for paragraph in document.paragraphs:
    if paragraph.text in normal_texts or paragraph.text.startswith("• "):
        paragraph.style = "Normal"
    if not paragraph.text.strip() and paragraph.style.name.startswith("Heading"):
        paragraph.style = "Normal"

# Correct the development section hierarchy.
for paragraph in document.paragraphs:
    if paragraph.text == "2.1 채널 데이터 모델":
        paragraph.style = "Heading 3"
    if paragraph.text == "9. API 실행 화면":
        paragraph.style = "Heading 2"

# Replace the empty section 7 placeholder with a concise final-state summary.
summary_anchor = next(p for p in document.paragraphs if p.text == "7. 차별화 포인트")
summary_anchor.text = "7. 최종 구현 및 검증 요약"
summary_anchor.style = "Heading 2"
future_anchor = next(p for p in document.paragraphs if p.text == "8. 향후 개선 사항")
summary_entries = [
    ("7.1 데이터 정합성과 동시성", "Heading 3"),
    ("• 동일 좌석 예매 100건: 성공 1건, 정상 거절 99건으로 중복 예매를 방지하였다.", "Normal"),
    ("• 재고 10개 상품 주문 100건: 성공 10건, 정상 거절 90건이며 최종 재고는 0이다.", "Normal"),
    ("• 동일 주문 취소 100건: 성공 1건, 정상 거절 99건이며 재고는 한 번만 복구된다.", "Normal"),
    ("• Gradle의 100% successful은 위 성공·거절 조건과 최종 DB 상태 검증이 모두 통과했다는 의미이다.", "Normal"),
    ("7.2 매출과 조회 최적화", "Heading 3"),
    ("• 예매 시점 좌석 가격을 reservedPrice로 보존하고 상품·공연·좌석 등급별 통계를 MyBatis로 제공한다.", "Normal"),
    ("• 취소 주문과 취소 예매는 집계에서 제외하며 날짜를 생략하면 요청 당일을 조회한다.", "Normal"),
    ("• 주문 항목·상품과 게시글 좋아요 수는 일괄 조회하여 N+1 문제를 줄였다.", "Normal"),
    ("7.3 보안과 오류 처리", "Heading 3"),
    ("• 역할과 담당 채널 관계를 함께 검증하며 공개 조회와 MANAGER 전용 API를 분리하였다.", "Normal"),
    ("• OpenAPI는 보호 API에만 bearerAuth를 표시하여 Postman에서도 공개 API가 No Auth로 구분된다.", "Normal"),
    ("• 예상하지 못한 오류는 내부 정보를 숨긴 INTERNAL_SERVER_ERROR 응답으로 변환하고 상세 내용은 서버 로그에 남긴다.", "Normal"),
    ("7.4 재현 가능한 테스트 데이터", "Heading 3"),
    ("• 모든 테이블에 10건 이상을 제공하고 INSERT ... ON CONFLICT DO NOTHING으로 중복 적재를 방지한다.", "Normal"),
    ("• 샘플 계정 30개는 비밀번호 1234로 로그인할 수 있으며 비밀번호는 BCrypt 해시로 저장한다.", "Normal"),
    ("• 채널마다 공연·상품 3개 이상과 VIP·R·S 좌석을 제공하고 2025~2026년 고정 날짜로 통계를 재현한다.", "Normal"),
]
for text, style in summary_entries:
    insert_before(future_anchor, text, style)

if CONCURRENCY_IMAGE.exists():
    image_paragraph = insert_before(future_anchor)
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.add_run().add_picture(str(CONCURRENCY_IMAGE), width=Inches(6.8))
    caption = insert_before(
        future_anchor,
        "그림 1. 100개 동시 요청 통합 테스트 결과(성공·정상 거절 건수)",
        "Caption",
    )
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Document page and header/footer layout.
for section in document.sections:
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)
    header = section.header.paragraphs[0]
    header.text = "BERRIES  |  팬 커뮤니티 플랫폼 백엔드"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = "맑은 고딕"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(100, 116, 139)
        set_east_asia_font(run._element, "맑은 고딕")
    footer = section.footer.paragraphs[0]
    footer.clear()
    add_page_number(footer)
    for run in footer.runs:
        run.font.name = "맑은 고딕"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(100, 116, 139)
        set_east_asia_font(run._element, "맑은 고딕")

# Global typography.
normal = document.styles["Normal"]
normal.font.name = "맑은 고딕"
normal.font.size = Pt(9.5)
normal.font.color.rgb = TEXT
set_east_asia_font(normal.element, "맑은 고딕")
normal.paragraph_format.space_after = Pt(4)
normal.paragraph_format.line_spacing = 1.15
normal.paragraph_format.widow_control = True

heading_specs = {
    "Heading 1": (18, NAVY, 14, 7),
    "Heading 2": (13.5, BLUE, 11, 5),
    "Heading 3": (11.5, TEAL, 8, 3),
}
for style_name, (size, color, before, after) in heading_specs.items():
    style = document.styles[style_name]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    set_east_asia_font(style.element, "맑은 고딕")
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

caption_style = document.styles["Caption"]
caption_style.font.name = "맑은 고딕"
caption_style.font.size = Pt(8.5)
caption_style.font.italic = False
caption_style.font.color.rgb = RGBColor(71, 85, 105)
set_east_asia_font(caption_style.element, "맑은 고딕")
caption_style.paragraph_format.space_before = Pt(3)
caption_style.paragraph_format.space_after = Pt(8)

code_style = ensure_style(document, "API Code")
code_style.font.name = "Consolas"
code_style.font.size = Pt(8.5)
code_style.font.color.rgb = RGBColor(30, 41, 59)
set_east_asia_font(code_style.element, "맑은 고딕")
code_style.paragraph_format.left_indent = Cm(0.35)
code_style.paragraph_format.right_indent = Cm(0.2)
code_style.paragraph_format.space_before = Pt(0)
code_style.paragraph_format.space_after = Pt(0)

# Paragraph-level styling and compact API code blocks.
http_methods = ("GET ", "POST ", "PATCH ", "PUT ", "DELETE ")
labels = {"Request", "Response", "Query", "Query Parameter", "요청 본문", "처리", "집계 규칙"}
for index, paragraph in enumerate(document.paragraphs):
    text = paragraph.text.strip()
    if index == 1 and text == "팬 커뮤니티 플랫폼 요구사항 정의서":
        paragraph.style = "Title"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if paragraph.style.name == "Heading 1":
        add_bottom_border(paragraph)
    if text.startswith(http_methods) or text in {"{", "}", "[", "]", "},"} or text.startswith('"'):
        paragraph.style = code_style
        p_pr = paragraph._p.get_or_add_pPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), LIGHT_GRAY)
        p_pr.append(shading)
    elif text in labels or text.startswith("권한:"):
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor.from_string(TEAL)
    if paragraph.style.name == "List Paragraph" and text:
        paragraph.style = "Caption"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Screenshot section labels.
domain_labels = {
    "전체 Swagger API", "Auth", "User", "Channel", "Channel-Sales", "Community",
    "Post", "Comment", "Product", "Order", "Concert", "Seat", "Reservation",
}
for paragraph in document.paragraphs:
    if paragraph.text.strip() in domain_labels:
        paragraph.style = "Heading 3"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

# Center images and keep their captions close.
for paragraph in document.paragraphs:
    if paragraph._p.xpath(".//w:drawing"):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.keep_with_next = True

# Consistent table styling.
for table in document.tables:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index == 0:
                set_cell_shading(cell, BLUE)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.name = "맑은 고딕"
                    run.font.size = Pt(8.5)
                    set_east_asia_font(run._element, "맑은 고딕")
                    if row_index == 0:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
        if row_index == 0:
            tr_pr = row._tr.get_or_add_trPr()
            repeat = OxmlElement("w:tblHeader")
            repeat.set(qn("w:val"), "true")
            tr_pr.append(repeat)

# Cover title styling.
title_style = document.styles["Title"]
title_style.font.name = "맑은 고딕"
title_style.font.size = Pt(26)
title_style.font.bold = True
title_style.font.color.rgb = RGBColor.from_string(NAVY)
set_east_asia_font(title_style.element, "맑은 고딕")
title_style.paragraph_format.space_before = Pt(24)
title_style.paragraph_format.space_after = Pt(18)

document.core_properties.title = "Berries 팬 커뮤니티 플랫폼 백엔드"
document.core_properties.subject = "요구사항, API 명세, 개발 내용 및 검증 결과"
document.core_properties.author = "SKALA JJAP BERRIES"
document.core_properties.keywords = "Spring Boot, JWT, MyBatis, JPA, 동시성, 팬 커뮤니티"

document.save(DOCUMENT_PATH)
