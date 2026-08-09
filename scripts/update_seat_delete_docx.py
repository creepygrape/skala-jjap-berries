from docx import Document
from docx.oxml import OxmlElement


DOCUMENT_PATH = "readme.docx"
MARKER = "좌석 개별 삭제"
POLICY = "운영 정책: 좌석 삭제는 공연 예매 시작 전까지만 허용한다."
UPDATE_POLICY = "운영 정책: 좌석 수정은 공연 예매 시작 전까지만 허용한다."
SALES_POLICY = "날짜 기본값: from과 to를 생략하면 API 요청 당일을 조회한다."
SEAT_QUERY_POLICY = "좌석 목록 조회는 MyBatis 동적 SQL로 필터링하며 page와 size로 페이징할 수 있다."
SEAT_PAGE_RESPONSE = (
    "응답은 content와 page, size, totalElements, totalPages, first, last "
    "페이지 메타데이터를 포함한다."
)
SALES_DETAIL_MARKER = "상품별 매출 상세: GET /api/analytics/sales/products?channelId={channelId}"
BULK_CREATE_POLICY = "좌석 일괄 등록의 startNumber는 선택값이다."
CONCERT_PERIOD_POLICY = "공연·좌석 운영 기간 정책"
REFACTORING_MARKER = "프로젝트 안정성 개선 사항"
API_MAPPING_MARKER = "도메인 기준 API 경로 규칙"
UNEXPECTED_ERROR_POLICY = "예상하지 못한 서버 오류 처리 정책"
SAMPLE_DATA_POLICY = "로컬 테스트용 초기 데이터"
OPENAPI_SECURITY_POLICY = "Swagger 인증 문서화 정책"
CHANNEL_MANAGER_QUERY_SECURITY_POLICY = "채널 매니저 조회 API 보안 정책"
ANALYTICS_SAMPLE_DATA_POLICY = "매출 집계 확인용 샘플 데이터"
FIXED_SAMPLE_DATE_POLICY = "초기 데이터 날짜 정책"
CONCURRENCY_TEST_POLICY = "100개 동시 요청 통합 테스트"
CONCURRENCY_RESULT_POLICY = "동시성 테스트 성공과 요청 거절 결과"


def insert_after(paragraph, text, style=None):
    new_xml = OxmlElement("w:p")
    paragraph._p.addnext(new_xml)
    new_paragraph = paragraph._parent.add_paragraph()
    new_paragraph._p.getparent().remove(new_paragraph._p)
    new_xml.getparent().replace(new_xml, new_paragraph._p)
    if style:
        new_paragraph.style = style
    elif paragraph.style:
        new_paragraph.style = paragraph.style
    new_paragraph.add_run(text)
    return new_paragraph


document = Document(DOCUMENT_PATH)
if not any(paragraph.text == MARKER for paragraph in document.paragraphs):
    anchor = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith("요구사항: AVAILABLE 좌석은 status를 RESERVED")
    )
    entries = [
        ("좌석 개별 삭제", "Heading 3"),
        ("DELETE /api/seats/{seatId}", None),
        ("담당 채널의 매니저가 좌석 ID로 좌석 하나를 삭제한다.", None),
        ("응답: 204 No Content", None),
        ("좌석 일괄 삭제", "Heading 3"),
        ("DELETE /api/concerts/{concertId}/seats/bulk", None),
        ("요청 본문", None),
        ('{\n  "section": "A",\n  "startNumber": 1,\n  "endNumber": 20\n}', None),
        (
            "요구사항: 시작 번호와 끝 번호를 모두 포함하는 연속 좌석을 한 번에 삭제한다. "
            "범위는 최대 500석이며, 범위 내 좌석이 하나라도 없으면 전체 요청을 실패한다.",
            None,
        ),
        (
            "삭제 제한: AVAILABLE 상태이고 예매 이력이 없는 좌석만 삭제할 수 있다. "
            "개별 삭제와 일괄 삭제 모두 담당 채널 매니저 권한이 필요하다.",
            None,
        ),
        ("응답: 204 No Content", None),
    ]
    for text, style in entries:
        anchor = insert_after(anchor, text, style)
    document.save(DOCUMENT_PATH)

document = Document(DOCUMENT_PATH)
if not any(paragraph.text == POLICY for paragraph in document.paragraphs):
    anchor = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith("삭제 제한: AVAILABLE 상태이고 예매 이력이 없는 좌석만")
    )
    insert_after(anchor, POLICY)
    document.save(DOCUMENT_PATH)

document = Document(DOCUMENT_PATH)
if not any(paragraph.text == UPDATE_POLICY for paragraph in document.paragraphs):
    anchor = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith("요구사항: AVAILABLE 좌석은 status를 RESERVED")
    )
    insert_after(anchor, UPDATE_POLICY)
    document.save(DOCUMENT_PATH)

document = Document(DOCUMENT_PATH)
if not any(paragraph.text == SALES_POLICY for paragraph in document.paragraphs):
    anchor = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text == "GET /api/analytics/{channelId}/sales"
    )
    entries = [
        (SALES_POLICY, None),
        ("선택 파라미터: from, to (yyyy-MM-dd, 각각 생략 시 요청 당일 적용)", None),
        (
            "매출 분류: 상품 주문은 totalOrderCount, totalQuantitySold, "
            "totalProductSalesAmount로 제공하고 공연 예매는 totalReservationCount, "
            "totalConcertSalesAmount로 제공한다.",
            None,
        ),
        ("totalSalesAmount는 상품 매출과 공연 매출을 합산한 금액이다.", None),
        ("dailySales에서도 상품과 공연 매출을 날짜별로 구분한다.", None),
    ]
    for text, style in entries:
        anchor = insert_after(anchor, text, style)
    document.save(DOCUMENT_PATH)

document = Document(DOCUMENT_PATH)
if not any(paragraph.text == SEAT_QUERY_POLICY for paragraph in document.paragraphs):
    anchor = next(paragraph for paragraph in document.paragraphs if paragraph.text == "좌석 목록 조회")
    entries = [
        (SEAT_QUERY_POLICY, None),
        ("선택 파라미터: section, grade, status, page(기본 0), size(기본 100, 최대 500)", None),
        ("응답 배열 형식은 기존과 동일하며 지정한 페이지의 좌석만 반환한다.", None),
    ]
    for text, style in entries:
        anchor = insert_after(anchor, text, style)
    document.save(DOCUMENT_PATH)

document = Document(DOCUMENT_PATH)
if not any(paragraph.text == SALES_DETAIL_MARKER for paragraph in document.paragraphs):
    anchor = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text in {
            "GET /api/analytics/{channelId}/sales",
            "GET /api/analytics/sales?channelId={channelId}",
        }
    )
    entries = [
        (SALES_DETAIL_MARKER, None),
        ("상품별 주문 건수, 판매 수량, 매출을 매출액 내림차순으로 반환한다.", None),
        ("공연별 매출 상세: GET /api/analytics/sales/concerts?channelId={channelId}", None),
        ("공연별 유효 예매 건수와 매출을 매출액 내림차순으로 반환한다.", None),
        ("좌석 등급별 판매율: GET /api/analytics/seat-grades?channelId={channelId}&concertId={concertId}", None),
        ("등급별 전체 좌석 수, 예매 좌석 수, 판매율(%), 유효 예매 매출을 반환한다.", None),
        ("취소된 주문과 취소된 예매는 모든 매출 상세 집계에서 제외한다.", None),
    ]
    for text, style in entries:
        anchor = insert_after(anchor, text, style)
    document.save(DOCUMENT_PATH)

document = Document(DOCUMENT_PATH)
if not any(paragraph.text == BULK_CREATE_POLICY for paragraph in document.paragraphs):
    anchor = next(paragraph for paragraph in document.paragraphs if paragraph.text == "좌석 일괄 등록")
    entries = [
        (BULK_CREATE_POLICY, None),
        (
            "생략 시 같은 공연·구역에 등록된 숫자 좌석 중 가장 큰 번호의 다음 번호부터 생성한다. "
            "해당 구역에 좌석이 없으면 1번부터 생성한다.",
            None,
        ),
        ("중간의 빈 번호는 자동으로 채우지 않으며 필요하면 startNumber를 직접 지정한다.", None),
        (
            "응답은 section, 실제 startNumber, endNumber, createdCount와 생성된 seats 배열을 포함한다.",
            None,
        ),
        ("자동 번호 계산과 등록은 공연 단위 비관적 잠금이 적용된 하나의 트랜잭션에서 처리한다.", None),
    ]
    for text, style in entries:
        anchor = insert_after(anchor, text, style)
    document.save(DOCUMENT_PATH)

document = Document(DOCUMENT_PATH)
if not any(paragraph.text == CONCERT_PERIOD_POLICY for paragraph in document.paragraphs):
    anchor = next(
        paragraph
        for index, paragraph in enumerate(document.paragraphs)
        if index > 500 and paragraph.text == "11. 공연"
    )
    entries = [
        (CONCERT_PERIOD_POLICY, "Heading 3"),
        ("공연 등록: 현재 시각 < 예매 시작 ≤ 예매 종료 ≤ 공연 일시 조건을 만족해야 한다.", None),
        ("예매 시작 전: 공연의 모든 정보와 좌석 등록·수정·삭제가 가능하다.", None),
        ("예매 시작 후: 공연 제목·장소 및 공연 상태만 변경할 수 있다.", None),
        ("예매 시작 후: 공연 일시와 예매 시작·종료 일시는 변경할 수 없다.", None),
        ("예매 시작 후: 좌석 개별·일괄 등록, 수정, 삭제는 모두 허용하지 않는다.", None),
        ("공연 삭제 API는 제공하지 않는다.", None),
    ]
    for text, style in entries:
        anchor = insert_after(anchor, text, style)
    document.save(DOCUMENT_PATH)

document = Document(DOCUMENT_PATH)
paragraphs = document.paragraphs
dev_start = next(
    (index for index, paragraph in enumerate(paragraphs) if paragraph.text == "Swagger 데이터 확인용 API"),
    None,
)
if dev_start is not None:
    dev_end = next(
        index
        for index in range(dev_start + 1, len(paragraphs))
        if paragraphs[index].text == "2. 인증 / 회원"
    )
    for paragraph in paragraphs[dev_start:dev_end]:
        paragraph._element.getparent().remove(paragraph._element)
    document.save(DOCUMENT_PATH)

document = Document(DOCUMENT_PATH)
for paragraph in list(document.paragraphs):
    if "Swagger 데이터 확인 API:" in paragraph.text:
        paragraph._element.getparent().remove(paragraph._element)
document.save(DOCUMENT_PATH)

document = Document(DOCUMENT_PATH)
if not any(paragraph.text == CONCURRENCY_RESULT_POLICY for paragraph in document.paragraphs):
    anchor = document.paragraphs[-1]
    entries = [
        (CONCURRENCY_RESULT_POLICY, "Heading 2"),
        ("Gradle 리포트의 100% successful은 세 테스트의 검증 조건이 모두 통과했다는 의미이며 100개 요청이 모두 성공했다는 의미가 아니다.", None),
        ("동일 좌석 예매: 총 100건 중 성공 1건, 정상 거절 99건을 검증한다.", None),
        ("재고 10개 상품 주문: 총 100건 중 성공 10건, 정상 거절 90건과 최종 재고 0을 검증한다.", None),
        ("동일 주문 취소: 총 100건 중 성공 1건, 정상 거절 99건과 재고 1회 복구를 검증한다.", None),
        ("정상 거절은 동시성 제어에 따른 BusinessException이며 예상하지 못한 예외가 발생하면 테스트가 실패한다.", None),
        ("각 테스트는 total, success, rejected 값을 Gradle 표준 출력 리포트에 기록한다.", None),
    ]
    for text, style in entries:
        anchor = insert_after(anchor, text, style)
    document.save(DOCUMENT_PATH)

document = Document(DOCUMENT_PATH)
if not any(paragraph.text == CONCURRENCY_TEST_POLICY for paragraph in document.paragraphs):
    anchor = document.paragraphs[-1]
    entries = [
        (CONCURRENCY_TEST_POLICY, "Heading 2"),
        ("동시성 검증은 HTTP 부하 도구가 아닌 Spring Boot 서비스 통합 테스트로 수행한다.", None),
        ("CountDownLatch로 100개 작업을 준비한 뒤 같은 시작 신호로 동시에 실행한다.", None),
        ("동일 좌석 예매 100건 중 1건만 성공하고 좌석과 예매 이력이 하나만 생성되는지 검증한다.", None),
        ("재고 10개 상품에 주문 100건이 들어오면 10건만 성공하고 최종 재고가 0인지 검증한다.", None),
        ("동일 주문 취소 100건 중 1건만 성공하고 상품 재고가 한 번만 복구되는지 검증한다.", None),
        ("100개 스레드와 DB 커넥션 대기를 고려해 준비 제한 15초, 작업 완료 제한 30초를 적용한다.", None),
    ]
    for text, style in entries:
        anchor = insert_after(anchor, text, style)
    document.save(DOCUMENT_PATH)

document = Document(DOCUMENT_PATH)
if not any(paragraph.text == FIXED_SAMPLE_DATE_POLICY for paragraph in document.paragraphs):
    anchor = document.paragraphs[-1]
    entries = [
        (FIXED_SAMPLE_DATE_POLICY, "Heading 2"),
        ("data.sql에서는 CURRENT_TIMESTAMP를 사용하지 않고 고정 TIMESTAMP 리터럴을 사용한다.", None),
        ("샘플 날짜는 2025년부터 2026년 사이, 1월부터 12월 사이, 매월 1일부터 25일 사이로 분산한다.", None),
        ("공연 일시와 예매 시작·종료 일시는 고정 날짜를 사용하면서 시간적 선후관계를 유지한다.", None),
        ("주문 및 예매 날짜도 고정하여 같은 기간으로 조회할 때 재현 가능한 매출 통계 결과를 제공한다.", None),
    ]
    for text, style in entries:
        anchor = insert_after(anchor, text, style)
    document.save(DOCUMENT_PATH)

document = Document(DOCUMENT_PATH)
if not any(
    paragraph.text == ANALYTICS_SAMPLE_DATA_POLICY for paragraph in document.paragraphs
):
    anchor = document.paragraphs[-1]
    entries = [
        (ANALYTICS_SAMPLE_DATA_POLICY, "Heading 2"),
        ("각 채널은 공연 3개와 상품 3개 이상을 가진다.", None),
        ("각 공연은 VIP, R, S의 세 가지 이상 좌석 등급을 가진다.", None),
        ("채널별 상품 주문과 공연 예매 일자를 서로 다르게 배치해 기간 조회 시 여러 일별 매출 결과가 반환된다.", None),
        ("각 채널에 복수 상품의 유효 주문과 복수 공연의 유효 예매를 제공해 상품별·공연별 집계를 확인할 수 있다.", None),
        ("취소 주문과 취소 예매 데이터도 유지해 매출 제외 조건을 함께 확인할 수 있다.", None),
    ]
    for text, style in entries:
        anchor = insert_after(anchor, text, style)
    document.save(DOCUMENT_PATH)

document = Document(DOCUMENT_PATH)
if not any(
    paragraph.text == CHANNEL_MANAGER_QUERY_SECURITY_POLICY
    for paragraph in document.paragraphs
):
    anchor = document.paragraphs[-1]
    entries = [
        (CHANNEL_MANAGER_QUERY_SECURITY_POLICY, "Heading 2"),
        ("공개 채널 조회는 GET /api/channels와 GET /api/channels/{channelId}만 허용한다.", None),
        ("GET /api/channels/{channelId}/members는 담당 채널 가입 사용자 정보이므로 MANAGER 인증이 필요하다.", None),
        ("GET /api/channels/managed도 로그인한 MANAGER 인증이 필요하다.", None),
        ("두 API는 Swagger에서 bearerAuth 잠금 대상으로 표시한다.", None),
    ]
    for text, style in entries:
        anchor = insert_after(anchor, text, style)
    document.save(DOCUMENT_PATH)

document = Document(DOCUMENT_PATH)
if not any(paragraph.text == OPENAPI_SECURITY_POLICY for paragraph in document.paragraphs):
    anchor = document.paragraphs[-1]
    entries = [
        (OPENAPI_SECURITY_POLICY, "Heading 2"),
        ("OpenAPI 최상위 전역 Bearer 인증 선언은 사용하지 않는다.", None),
        ("회원가입·로그인·토큰 재발급과 공개 GET API에는 보안 요구사항을 선언하지 않는다.", None),
        ("인증이 필요한 API 작업에만 bearerAuth 보안 요구사항을 선언한다.", None),
        ("Postman으로 OpenAPI 문서를 가져오면 공개 API는 No Auth, 보호 API는 Bearer Token 대상으로 구분된다.", None),
        ("Spring Security 접근 정책과 JWT 필터의 토큰 검증 동작은 변경하지 않는다.", None),
    ]
    for text, style in entries:
        anchor = insert_after(anchor, text, style)
    document.save(DOCUMENT_PATH)

document = Document(DOCUMENT_PATH)
if not any(paragraph.text == SAMPLE_DATA_POLICY for paragraph in document.paragraphs):
    anchor = document.paragraphs[-1]
    entries = [
        (SAMPLE_DATA_POLICY, "Heading 2"),
        ("src/main/resources/data.sql은 H2 로컬 테스트용 샘플 데이터를 초기화한다.", None),
        ("모든 테이블에 최소 10건을 제공하며 고정 ID 1001 이상과 INSERT ... ON CONFLICT DO NOTHING을 사용해 재실행 시 중복을 방지한다.", None),
        ("사용자: MANAGER 10명(manager01@berries.test~manager10@berries.test), ARTIST 10명(artist01@berries.test~artist10@berries.test), USER 10명(user01@berries.test~user10@berries.test)", None),
        ("모든 샘플 사용자의 로그인 비밀번호는 1234이며 DB에는 BCrypt 해시만 저장한다.", None),
        ("채널별 담당 매니저와 아티스트, 팬 멤버십을 연결해 역할 및 권한 API를 테스트할 수 있다.", None),
        ("공연은 예매 진행 중과 예매 시작 전 상태를, 상품은 판매 중·판매 중지·품절 상태를 포함한다.", None),
        ("예매·취소 예매, 정상·취소 주문, 가격 스냅샷과 날짜별 매출 데이터가 포함되어 통계 API를 테스트할 수 있다.", None),
        ("게시글은 FAN·ARTIST·NOTICE 유형과 소프트 삭제 예시를, 댓글은 최상위 댓글·답글·소프트 삭제 예시를 포함한다.", None),
        ("refresh_token과 revoked_access_token 값은 저장·만료 정리 테스트용이며 실제 인증에 사용할 수 없는 샘플 문자열이다.", None),
    ]
    for text, style in entries:
        anchor = insert_after(anchor, text, style)
    document.save(DOCUMENT_PATH)

document = Document(DOCUMENT_PATH)
if not any(paragraph.text == UNEXPECTED_ERROR_POLICY for paragraph in document.paragraphs):
    anchor = document.paragraphs[-1]
    entries = [
        (UNEXPECTED_ERROR_POLICY, "Heading 2"),
        ("처리되지 않은 예외는 공통 예외 처리기가 HTTP 500 응답으로 변환한다.", None),
        ("응답 코드는 INTERNAL_SERVER_ERROR이며 메시지는 '서버 내부 오류가 발생했습니다.'로 고정한다.", None),
        ("예외 상세 내용과 스택 트레이스는 서버 로그에만 기록하고 API 응답에는 노출하지 않는다.", None),
        ("서버 로그에는 요청 HTTP 메서드와 URI를 함께 기록해 장애 추적에 사용한다.", None),
    ]
    for text, style in entries:
        anchor = insert_after(anchor, text, style)
    document.save(DOCUMENT_PATH)

document = Document(DOCUMENT_PATH)
path_replacements = {
    "/api/analytics/{channelId}/concerts/{concertId}/seat-grades": "/api/analytics/seat-grades?channelId={channelId}&concertId={concertId}",
    "/api/analytics/{channelId}/sales/products": "/api/analytics/sales/products?channelId={channelId}",
    "/api/analytics/{channelId}/sales/concerts": "/api/analytics/sales/concerts?channelId={channelId}",
    "/api/analytics/{channelId}/sales": "/api/analytics/sales?channelId={channelId}",
    "/api/concerts/{concertId}/seats/bulk": "/api/seats/bulk?concertId={concertId}",
    "/api/concerts/{concertId}/seats": "/api/seats?concertId={concertId}",
    "/api/concerts/{concertId}/reservations": "/api/reservations?concertId={concertId}",
    "/api/channels/{channelId}/concerts": "/api/concerts?channelId={channelId}",
    "/api/channels/{channelId}/products": "/api/products?channelId={channelId}",
    "/api/channels/{channelId}/posts": "/api/posts?channelId={channelId}",
    "/api/channels/{channelId}/memberships": "/api/memberships?channelId={channelId}",
    "/api/users/me/memberships": "/api/memberships/me",
    "/api/posts/{postId}/comments": "/api/comments?postId={postId}",
    "/api/managers/channels/{channelId}/managers/{managerId}": "/api/channels/{channelId}/managers/{managerId}",
    "/api/managers/channels/{channelId}/managers": "/api/channels/{channelId}/managers",
    "/api/managers/channels/{channelId}/members": "/api/channels/{channelId}/members",
    "/api/managers/channels/members": "/api/channels/{channelId}/members",
    "/api/channels/members": "/api/channels/{channelId}/members",
    "/api/managers/me/channels": "/api/channels/managed",
    "/api/managers/channels/{channelId}": "/api/channels/{channelId}",
    "/api/managers/channels": "/api/channels",
}
for paragraph in document.paragraphs:
    updated = paragraph.text
    for old_path, new_path in path_replacements.items():
        updated = updated.replace(old_path, new_path)
    if updated != paragraph.text:
        paragraph.text = updated
document.save(DOCUMENT_PATH)

document = Document(DOCUMENT_PATH)
for paragraph in document.paragraphs:
    if "고정 ID 1001 이상과 MERGE 문을 사용해" in paragraph.text:
        paragraph.text = paragraph.text.replace(
            "고정 ID 1001 이상과 MERGE 문을 사용해",
            "고정 ID 1001 이상과 INSERT ... ON CONFLICT DO NOTHING을 사용해",
        )
document.save(DOCUMENT_PATH)

document = Document(DOCUMENT_PATH)
paragraphs = list(document.paragraphs)
for index, paragraph in enumerate(paragraphs):
    if paragraph.text == "응답 배열 형식은 기존과 동일하며 지정한 페이지의 좌석만 반환한다.":
        paragraph.text = SEAT_PAGE_RESPONSE
    if '“channelId”: 1,' in paragraph.text:
        previous_texts = [candidate.text for candidate in paragraphs[max(0, index - 6):index]]
        if any("/api/channels/{channelId}/members" in text for text in previous_texts):
            paragraph._element.getparent().remove(paragraph._element)
document.save(DOCUMENT_PATH)

document = Document(DOCUMENT_PATH)
if not any(paragraph.text == API_MAPPING_MARKER for paragraph in document.paragraphs):
    anchor = document.paragraphs[-1]
    entries = [
        (API_MAPPING_MARKER, "Heading 2"),
        ("각 컨트롤러는 /api/{domain} 기본 경로를 사용하며 메서드에는 하위 경로만 선언한다.", None),
        ("권한 역할을 URL에 표시하던 /api/managers 경로는 제거하고 서비스 권한 검증으로 처리한다.", None),
        ("소속 채널이나 공연은 channelId, concertId 쿼리 파라미터로 전달한다.", None),
        ("ConcertController와 SeatController, PostController와 CommentController를 각각 분리하였다.", None),
    ]
    for text, style in entries:
        anchor = insert_after(anchor, text, style)
    document.save(DOCUMENT_PATH)

document = Document(DOCUMENT_PATH)
for paragraph in document.paragraphs:
    if '"seatNumber": "12"' in paragraph.text:
        paragraph.text = paragraph.text.replace(
            '"seatNumber": "12"', '"seatSequence": 12,\n    "seatLabel": "A-12"'
        )
    if "application-dev.yml은 개발용 토큰 만료 및 로그 설정에 사용한다" in paragraph.text:
        paragraph.text = paragraph.text.replace(
            "application-dev.yml은 개발용 토큰 만료 및 로그 설정에 사용한다",
            "application-dev.yml은 개발용 토큰 만료 설정에만 사용한다",
        )
    if "Request Body로 channelId와 ARTIST 역할 회원의 artistId를 받는다" in paragraph.text:
        paragraph.text = paragraph.text.replace(
            "Request Body로 channelId와 ARTIST 역할 회원의 artistId를 받는다",
            "경로로 channelId를 받고 Request Body에는 ARTIST 역할 회원의 artistId만 받는다",
        )
document.save(DOCUMENT_PATH)

document = Document(DOCUMENT_PATH)
if not any(paragraph.text == REFACTORING_MARKER for paragraph in document.paragraphs):
    anchor = document.paragraphs[-1]
    entries = [
        (REFACTORING_MARKER, "Heading 2"),
        ("개발 데이터 조회용 /api/dev API와 관련 Swagger 명세를 제거하였다.", None),
        ("공통 application.yml에 H2·JPA·로그 설정을 통합하고 application-dev.yml은 개발용 토큰 만료 설정에만 사용한다.", None),
        ("주문 취소는 주문과 상품을 ID 순서로 비관적 잠금하여 재고가 한 번만 복원되도록 한다.", None),
        ("상품 수정과 상태 변경에도 비관적 잠금을 적용해 주문의 재고 차감과 경합하지 않도록 한다.", None),
        ("주문 목록은 주문 항목과 상품을 한 번에 조회해 N+1 조회를 제거하였다.", None),
        ("예매 시점의 좌석 가격을 reservedPrice로 저장하며 매출 통계도 해당 가격 스냅샷을 사용한다.", None),
        ("매니저는 자기 자신의 채널 관리 권한을 제거할 수 없다.", None),
        ("DB 무결성 위반은 DATA_INTEGRITY_CONFLICT 공통 오류 응답으로 변환한다.", None),
        ("좌석은 section, seatSequence, seatLabel로 관리하며 범위·정렬·자동 번호는 seatSequence를 사용한다.", None),
        ("기존 H2 좌석 데이터 전환은 scripts/migrate_h2_seat_schema.sql을 사용한다.", None),
        ("만료 토큰 자동 정리는 현재 적용하지 않으며 추후 운영 DB의 정리 작업 또는 외부 스케줄 작업으로 결정한다.", None),
        ("좌석 조회 page는 최대 1,000,000이며 offset은 long 타입으로 계산한다.", None),
    ]
    for text, style in entries:
        anchor = insert_after(anchor, text, style)
document.save(DOCUMENT_PATH)

document = Document(DOCUMENT_PATH)
for paragraph in document.paragraphs:
    if "만료된 Refresh Token과 폐기 Access Token은 매시간 Spring 스케줄러로 정리한다." in paragraph.text:
        paragraph.text = (
            "만료 토큰 자동 정리는 현재 적용하지 않으며 추후 운영 DB의 정리 작업 또는 "
            "외부 스케줄 작업으로 결정한다."
        )
document.save(DOCUMENT_PATH)
